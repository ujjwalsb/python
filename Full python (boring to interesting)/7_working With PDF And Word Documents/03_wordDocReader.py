import docx			# sudo python3 -m pip install -U python-docx

doc = docx.Document('demo.docx')
print(len(doc.paragraphs))

print(doc.paragraphs[0].text)
